import os
import json
import uuid
from typing import List, Dict, Any, Optional
from pathlib import Path
from datetime import datetime

import faiss
import numpy as np
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from sentence_transformers import SentenceTransformer
from groq import Groq
from dotenv import load_dotenv
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

# Load environment variables from .env file
load_dotenv()

app = FastAPI(title="Question Paper Generator API", version="1.0.0")

# CORS middleware for frontend communication
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],  # Vite default port
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global variables for model and index
model = None
index = None
documents = []
index_path = Path("faiss_index")
uploaded_notes_path = Path("uploaded_notes")

# Create directories if they don't exist
index_path.mkdir(exist_ok=True)
uploaded_notes_path.mkdir(exist_ok=True)

# Fixed question categories and their rules
QUESTION_CATEGORIES = {
    "mcq": {
        "name": "Multiple Choice Questions",
        "description": "Questions with 4 options (A, B, C, D)",
        "default_count": 5,
        "marks_per_question": 1
    },
    "short_answer": {
        "name": "Short Answer Questions",
        "description": "Brief answers (2-3 sentences)",
        "default_count": 3,
        "marks_per_question": 2
    },
    "long_answer": {
        "name": "Long Answer Questions",
        "description": "Detailed answers (1-2 paragraphs)",
        "default_count": 2,
        "marks_per_question": 5
    },
    "case_study": {
        "name": "Case Study Questions",
        "description": "Scenario-based questions with multiple parts",
        "default_count": 1,
        "marks_per_question": 10
    },
    "application_based": {
        "name": "Application-Based Questions",
        "description": "Questions requiring practical application of concepts",
        "default_count": 2,
        "marks_per_question": 3
    },
    "numerical": {
        "name": "Numerical Questions",
        "description": "Questions requiring calculations and numerical answers",
        "default_count": 2,
        "marks_per_question": 4
    }
}

# Pydantic models

class Question(BaseModel):
    id: str
    type: str
    question: str
    options: Optional[List[str]] = None
    correct_answer: Optional[str] = None
    marks: int
    difficulty: str

class QuestionPaper(BaseModel):
    id: str
    subject: str
    total_marks: int
    duration: int
    difficulty: str
    generated_at: str
    questions: List[Question]

class UploadResponse(BaseModel):
    message: str
    session_id: str
    files_uploaded: List[str]

def load_notes(session_id: str = None) -> List[Dict[str, str]]:
    """Load all text files from uploaded notes or default notes directory."""
    documents = []
    
    if session_id:
        # Load from uploaded notes for specific session
        session_path = uploaded_notes_path / session_id
        if session_path.exists():
            for file_path in session_path.glob("*.txt"):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read().strip()
                    if content:
                        documents.append({
                            "content": content,
                            "filename": file_path.name,
                            "session_id": session_id
                        })
    else:
        # Load from default notes directory
        notes_path = Path("../notes")
        if notes_path.exists():
            for file_path in notes_path.glob("*.txt"):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read().strip()
                    if content:
                        documents.append({
                            "content": content,
                            "filename": file_path.name,
                            "session_id": "default"
                        })
    
    return documents

def create_embeddings(texts: List[str]) -> np.ndarray:
    """Create embeddings for the given texts."""
    global model
    if model is None:
        model = SentenceTransformer('all-MiniLM-L6-v2')
    
    embeddings = model.encode(texts)
    return embeddings

def build_faiss_index(session_id: str = None):
    """Build or rebuild the FAISS index from notes."""
    global index, documents
    
    print(f"Loading notes for session: {session_id}")
    documents = load_notes(session_id)
    
    if not documents:
        print("No notes found.")
        return
    
    print(f"Found {len(documents)} notes")
    
    # Create embeddings
    texts = [doc["content"] for doc in documents]
    embeddings = create_embeddings(texts)
    
    # Create FAISS index
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatIP(dimension)  # Inner product for cosine similarity
    
    # Normalize embeddings for cosine similarity
    faiss.normalize_L2(embeddings)
    index.add(embeddings)
    
    # Save index
    index_file = index_path / f"index_{session_id or 'default'}.faiss"
    docs_file = index_path / f"documents_{session_id or 'default'}.json"
    
    faiss.write_index(index, str(index_file))
    
    # Save document metadata
    with open(docs_file, 'w') as f:
        json.dump(documents, f, indent=2)
    
    print(f"FAISS index built with {index.ntotal} documents")

def load_faiss_index(session_id: str = None):
    """Load existing FAISS index if available."""
    global index, documents
    
    index_file = index_path / f"index_{session_id or 'default'}.faiss"
    docs_file = index_path / f"documents_{session_id or 'default'}.json"
    
    if index_file.exists() and docs_file.exists():
        try:
            index = faiss.read_index(str(index_file))
            with open(docs_file, 'r') as f:
                documents = json.load(f)
            print(f"Loaded FAISS index with {index.ntotal} documents")
        except Exception as e:
            print(f"Error loading index: {e}")
            build_faiss_index(session_id)
    else:
        build_faiss_index(session_id)

def retrieve_relevant_docs(query: str, k: int = 5) -> List[Dict[str, Any]]:
    """Retrieve relevant documents for the query."""
    global model, index, documents
    
    if index is None or not documents:
        return []
    
    # Create query embedding
    query_embedding = create_embeddings([query])
    faiss.normalize_L2(query_embedding)
    
    # Search
    scores, indices = index.search(query_embedding, k)
    
    relevant_docs = []
    for score, idx in zip(scores[0], indices[0]):
        if idx < len(documents):
            relevant_docs.append({
                "content": documents[idx]["content"],
                "filename": documents[idx]["filename"],
                "score": float(score)
            })
    
    return relevant_docs

def get_groq_client():
    """Get Groq client with API key."""
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        raise ValueError("GROQ_API_KEY not found in environment variables")
    return Groq(api_key=api_key)

def generate_question(question_type: str, context: str, difficulty: str = "medium") -> Dict[str, Any]:
    """Generate a question of specified type using Groq API."""
    
    print(f"Generating {question_type} question with context length: {len(context)}")
    print(f"GROQ_API_KEY is set: {bool(os.getenv('GROQ_API_KEY'))}")
    
    # Hardcoded prompts for each question type
    prompts = {
        "mcq": f"""Generate a multiple choice question based on the following context. The question should have exactly 4 options (A, B, C, D) with only one correct answer. Make it {difficulty} difficulty.

Context: {context}

Format your response as JSON:
{{
    "question": "The question text here",
    "options": ["Option A", "Option B", "Option C", "Option D"],
    "correct_answer": "A"
}}""",

        "short_answer": f"""Generate a short answer question based on the following context. The answer should be 2-3 sentences. Make it {difficulty} difficulty.

Context: {context}

Format your response as JSON:
{{
    "question": "The question text here",
    "correct_answer": "The expected answer in 2-3 sentences"
}}""",

        "long_answer": f"""Generate a long answer question based on the following context. The answer should be 1-2 paragraphs with detailed explanation. Make it {difficulty} difficulty.

Context: {context}

Format your response as JSON:
{{
    "question": "The question text here",
    "correct_answer": "The expected detailed answer in 1-2 paragraphs"
}}""",

        "case_study": f"""Generate a case study question based on the following context. Include a scenario and ask multiple related questions. Make it {difficulty} difficulty.

Context: {context}

Format your response as JSON:
{{
    "question": "Case study scenario and questions here",
    "correct_answer": "Expected answers for all parts"
}}""",

        "application_based": f"""Generate an application-based question that requires practical application of concepts from the context. Make it {difficulty} difficulty.

Context: {context}

Format your response as JSON:
{{
    "question": "The application-based question here",
    "correct_answer": "The expected answer showing practical application"
}}""",

        "numerical": f"""Generate a numerical question based on the following context. Include specific numbers and require calculations. Make it {difficulty} difficulty.

Context: {context}

Format your response as JSON:
{{
    "question": "The numerical question with given data here",
    "correct_answer": "The numerical answer with steps"
}}"""
    }

    try:
        client = get_groq_client()
        
        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {"role": "system", "content": "You are an expert question paper generator. Generate high-quality educational questions based on the provided context. Always respond with valid JSON format."},
                {"role": "user", "content": prompts[question_type]}
            ],
            max_tokens=500,
            temperature=0.7
        )
        
        answer_text = response.choices[0].message.content
        
        # Parse JSON response
        try:
            question_data = json.loads(answer_text)
            return question_data
        except json.JSONDecodeError:
            # Fallback if JSON parsing fails
            return {
                "question": f"Sample {question_type} question based on the context.",
                "options": ["Option A", "Option B", "Option C", "Option D"] if question_type == "mcq" else None,
                "correct_answer": "Sample answer"
            }
        
    except Exception as e:
        print(f"Groq API error: {e}")
        print(f"Error type: {type(e).__name__}")
        print(f"GROQ_API_KEY is set: {bool(os.getenv('GROQ_API_KEY'))}")
        print(f"GROQ_API_KEY value: {os.getenv('GROQ_API_KEY', 'NOT_SET')[:10]}..." if os.getenv('GROQ_API_KEY') else "NOT_SET")
        # Fallback response
        return {
            "question": f"Sample {question_type} question: Explain the key concepts from the provided context.",
            "options": ["Option A", "Option B", "Option C", "Option D"] if question_type == "mcq" else None,
            "correct_answer": "Sample answer based on the context provided."
        }

def export_to_docx(question_paper: QuestionPaper, file_path: str):
    """Export question paper to DOCX format."""
    doc = Document()
    
    # Title
    title = doc.add_heading(f'{question_paper.subject} - Question Paper', 0)
    title.alignment = 1  # Center alignment
    
    # Paper details
    doc.add_paragraph(f'Total Marks: {question_paper.total_marks}')
    doc.add_paragraph(f'Duration: {question_paper.duration} minutes')
    doc.add_paragraph(f'Difficulty: {question_paper.difficulty.title()}')
    doc.add_paragraph(f'Generated on: {question_paper.generated_at}')
    
    doc.add_paragraph()  # Empty line
    
    # Instructions
    doc.add_heading('Instructions:', level=1)
    doc.add_paragraph('• Read all questions carefully before attempting.')
    doc.add_paragraph('• Answer all questions in the order given.')
    doc.add_paragraph('• Show all working for numerical questions.')
    doc.add_paragraph('• Write clearly and legibly.')
    
    doc.add_paragraph()  # Empty line
    
    # Questions by category
    current_category = None
    question_number = 1
    
    for question in question_paper.questions:
        if question.type != current_category:
            current_category = question.type
            category_info = QUESTION_CATEGORIES.get(current_category, {})
            doc.add_heading(f'{category_info.get("name", current_category.title())} ({question.marks} marks each)', level=1)
        
        # Question text
        doc.add_paragraph(f'Q{question_number}. {question.question}')
        
        # Options for MCQ
        if question.options:
            for i, option in enumerate(question.options):
                doc.add_paragraph(f'   {chr(65+i)}. {option}', style='List Bullet')
        
        doc.add_paragraph()  # Empty line
        question_number += 1
    
    doc.save(file_path)

def export_to_pdf(question_paper: QuestionPaper, file_path: str):
    """Export question paper to PDF format."""
    doc = SimpleDocTemplate(file_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1  # Center
    )
    story.append(Paragraph(f'{question_paper.subject} - Question Paper', title_style))
    story.append(Spacer(1, 12))
    
    # Paper details
    story.append(Paragraph(f'<b>Total Marks:</b> {question_paper.total_marks}', styles['Normal']))
    story.append(Paragraph(f'<b>Duration:</b> {question_paper.duration} minutes', styles['Normal']))
    story.append(Paragraph(f'<b>Difficulty:</b> {question_paper.difficulty.title()}', styles['Normal']))
    story.append(Paragraph(f'<b>Generated on:</b> {question_paper.generated_at}', styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Instructions
    story.append(Paragraph('<b>Instructions:</b>', styles['Heading2']))
    story.append(Paragraph('• Read all questions carefully before attempting.', styles['Normal']))
    story.append(Paragraph('• Answer all questions in the order given.', styles['Normal']))
    story.append(Paragraph('• Show all working for numerical questions.', styles['Normal']))
    story.append(Paragraph('• Write clearly and legibly.', styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Questions by category
    current_category = None
    question_number = 1
    
    for question in question_paper.questions:
        if question.type != current_category:
            current_category = question.type
            category_info = QUESTION_CATEGORIES.get(current_category, {})
            story.append(Paragraph(f'<b>{category_info.get("name", current_category.title())} ({question.marks} marks each)</b>', styles['Heading2']))
        
        # Question text
        story.append(Paragraph(f'Q{question_number}. {question.question}', styles['Normal']))
        
        # Options for MCQ
        if question.options:
            for i, option in enumerate(question.options):
                story.append(Paragraph(f'   {chr(65+i)}. {option}', styles['Normal']))
        
        story.append(Spacer(1, 12))
        question_number += 1
    
    doc.build(story)

@app.on_event("startup")
async def startup_event():
    """Initialize the application on startup."""
    print("Starting Question Paper Generator API...")
    load_faiss_index()

@app.post("/upload_notes", response_model=UploadResponse)
async def upload_notes(files: List[UploadFile] = File(...)):
    """Upload unit notes for question paper generation."""
    session_id = str(uuid.uuid4())
    session_path = uploaded_notes_path / session_id
    session_path.mkdir(exist_ok=True)
    
    uploaded_files = []
    
    for file in files:
        if not file.filename.endswith('.txt'):
            raise HTTPException(status_code=400, detail="Only .txt files are supported")
        
        file_path = session_path / file.filename
        content = await file.read()
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content.decode('utf-8'))
        
        uploaded_files.append(file.filename)
    
    # Build index for this session
    build_faiss_index(session_id)
    
    return UploadResponse(
        message=f"Successfully uploaded {len(uploaded_files)} files",
        session_id=session_id,
        files_uploaded=uploaded_files
    )

@app.post("/generate_paper", response_model=QuestionPaper)
async def generate_paper(
    categories: str = Form(...),
    subject: str = Form("General"),
    duration: int = Form(60),
    difficulty: str = Form("medium"),
    total_marks: int = Form(0),
    session_id: str = Form(None)
):
    """Generate a question paper based on requirements."""
    
    # Parse categories from JSON string
    try:
        categories_dict = json.loads(categories)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid categories format")
    
    # Load index for the session
    load_faiss_index(session_id)
    
    if not documents:
        raise HTTPException(status_code=400, detail="No notes available. Please upload notes first.")
    
    # Validate categories
    for category in categories_dict:
        if category not in QUESTION_CATEGORIES:
            raise HTTPException(status_code=400, detail=f"Invalid category: {category}")
    
    # Generate questions
    questions = []
    question_id = 1
    
    for category, count in categories_dict.items():
        if count <= 0:
            continue
            
        category_info = QUESTION_CATEGORIES[category]
        
        for i in range(count):
            # Get relevant context for this question
            relevant_docs = retrieve_relevant_docs(f"{category} {category_info['name']}", k=3)
            context = "\n\n".join([doc["content"] for doc in relevant_docs])
            
            if not context:
                context = "\n\n".join([doc["content"] for doc in documents[:2]])
            
            # Generate question
            question_data = generate_question(category, context, difficulty)
            
            question = Question(
                id=str(question_id),
                type=category,
                question=question_data["question"],
                options=question_data.get("options"),
                correct_answer=question_data.get("correct_answer"),
                marks=category_info["marks_per_question"],
                difficulty=difficulty
            )
            
            questions.append(question)
            question_id += 1
    
    # Calculate total marks
    calculated_total_marks = sum(q.marks for q in questions)
    
    # Create question paper
    paper_id = str(uuid.uuid4())
    question_paper = QuestionPaper(
        id=paper_id,
        subject=subject,
        total_marks=calculated_total_marks,
        duration=duration,
        difficulty=difficulty,
        generated_at=datetime.now().isoformat(),
        questions=questions
    )
    
    # Save to file
    paper_file = index_path / f"paper_{paper_id}.json"
    with open(paper_file, 'w') as f:
        json.dump(question_paper.dict(), f, indent=2)
    
    return question_paper

@app.get("/categories")
async def get_categories():
    """Get available question categories."""
    return QUESTION_CATEGORIES

@app.get("/download/{paper_id}/{format}")
async def download_paper(paper_id: str, format: str):
    """Download question paper in specified format."""
    paper_file = index_path / f"paper_{paper_id}.json"
    
    if not paper_file.exists():
        raise HTTPException(status_code=404, detail="Question paper not found")
    
    with open(paper_file, 'r') as f:
        paper_data = json.load(f)
    
    question_paper = QuestionPaper(**paper_data)
    
    if format == "json":
        return FileResponse(paper_file, filename=f"question_paper_{paper_id}.json")
    elif format == "docx":
        docx_file = index_path / f"paper_{paper_id}.docx"
        export_to_docx(question_paper, str(docx_file))
        return FileResponse(docx_file, filename=f"question_paper_{paper_id}.docx")
    elif format == "pdf":
        pdf_file = index_path / f"paper_{paper_id}.pdf"
        export_to_pdf(question_paper, str(pdf_file))
        return FileResponse(pdf_file, filename=f"question_paper_{paper_id}.pdf")
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Supported: json, docx, pdf")

@app.get("/health")
async def health_check():
    """Health check endpoint."""
    return {
        "status": "healthy",
        "indexed_documents": len(documents) if documents else 0,
        "index_loaded": index is not None,
        "available_categories": list(QUESTION_CATEGORIES.keys())
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)